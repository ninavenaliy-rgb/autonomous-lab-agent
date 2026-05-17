"""
DOCX Report Builder.
Constructs a fully formatted laboratory report with:
- ГОСТ-style formatting (Times New Roman 14pt, 1.5 line spacing)
- Sequential figure and table captions
- Automatic table of contents field
- Embedded screenshots
- Section structure matching lab methodology
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from loguru import logger

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Cm, Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.error("python-docx not installed — report generation disabled")

from core.config import get_config
from report.caption_manager import CaptionManager, CaptionEntry, get_caption_manager
from report.toc_manager import TOCManager


@dataclass
class ReportSection:
    title: str
    level: int = 1          # 1=chapter, 2=section, 3=subsection
    content: str = ""
    figures: list[str] = field(default_factory=list)  # image file paths
    figure_captions: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)  # rows x cols x text
    table_captions: list[str] = field(default_factory=list)
    task_id: str = ""


@dataclass
class ReportMeta:
    title: str
    subject: str = ""
    lab_number: str = ""
    student: str = ""
    group: str = ""
    teacher: str = ""
    university: str = ""
    department: str = ""
    year: str = ""
    city: str = ""


class DocxBuilder:
    """
    Builds a complete DOCX lab report.
    Applies ГОСТ 7.32-2017 and university-standard formatting.
    """

    def __init__(self) -> None:
        cfg = get_config()
        self._cfg = cfg.report
        self._captions = get_caption_manager()
        self._toc = TOCManager()
        self._doc: Any = None

    def new_document(self) -> None:
        """Start a fresh document with all base styles configured."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        self._doc = Document()
        self._setup_page_margins()
        self._setup_styles()
        self._captions.reset()

    def _setup_page_margins(self) -> None:
        sections = self._doc.sections
        for section in sections:
            section.top_margin = Cm(self._cfg.page_margin_cm * 1.0)
            section.bottom_margin = Cm(self._cfg.page_margin_cm * 1.0)
            section.left_margin = Cm(3.0)   # ГОСТ: 3cm left
            section.right_margin = Cm(1.5)  # ГОСТ: 1.5cm right

    def _setup_styles(self) -> None:
        """Configure base document styles to ГОСТ standard."""
        doc = self._doc
        styles = doc.styles

        # Normal paragraph style
        normal = styles["Normal"]
        normal.font.name = self._cfg.font_name
        normal.font.size = Pt(self._cfg.font_size_body)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = Pt(self._cfg.font_size_body * self._cfg.line_spacing)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(1.25)  # ГОСТ: 1.25cm paragraph indent

        # Heading 1
        try:
            h1 = styles["Heading 1"]
            h1.font.name = self._cfg.font_name
            h1.font.size = Pt(self._cfg.font_size_heading1)
            h1.font.bold = True
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(6)
            h1.paragraph_format.first_line_indent = Cm(0)
        except KeyError:
            pass

        # Heading 2
        try:
            h2 = styles["Heading 2"]
            h2.font.name = self._cfg.font_name
            h2.font.size = Pt(self._cfg.font_size_heading2)
            h2.font.bold = True
            h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h2.paragraph_format.first_line_indent = Cm(0)
        except KeyError:
            pass

    def add_title_page(self, meta: ReportMeta) -> None:
        """Add a formatted title page."""
        doc = self._doc

        def centered(text: str, size: int = 14, bold: bool = False, before: int = 0) -> Any:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(before)
            run = p.add_run(text)
            run.font.name = self._cfg.font_name
            run.font.size = Pt(size)
            run.font.bold = bold
            return p

        if meta.university:
            centered(meta.university.upper(), size=12, bold=False, before=0)
        if meta.department:
            centered(meta.department, size=12)

        centered("", before=60)
        centered("ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ", size=16, bold=True)
        if meta.lab_number:
            centered(f"№ {meta.lab_number}", size=14, bold=True)

        centered("", before=12)
        centered(meta.title, size=14, bold=True)

        if meta.subject:
            centered("", before=12)
            centered(f"По дисциплине: {meta.subject}", size=12)

        # Student info — right-aligned
        centered("", before=48)
        for label, val in [
            ("Выполнил:", meta.student),
            ("Группа:", meta.group),
            ("Проверил:", meta.teacher),
        ]:
            if val:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(f"{label} {val}")
                run.font.name = self._cfg.font_name
                run.font.size = Pt(12)

        centered("", before=48)
        parts = []
        if meta.city:
            parts.append(meta.city)
        if meta.year:
            parts.append(meta.year)
        if parts:
            centered(" — ".join(parts), size=12)

        doc.add_page_break()

    def add_toc(self) -> None:
        """Insert the automatic Table of Contents."""
        self._toc.insert_toc_field(self._doc)

    def add_section(self, section: ReportSection) -> None:
        """Add a complete report section with text, figures, and tables."""
        doc = self._doc

        # Heading
        style = f"Heading {section.level}"
        try:
            heading = doc.add_heading(section.title, level=section.level)
            heading.style.font.name = self._cfg.font_name
        except Exception:
            p = doc.add_paragraph(section.title)
            p.style = doc.styles.get("Normal") or doc.styles["Normal"]

        self._toc.register_heading(section.title, section.level)

        # Body text
        if section.content:
            for paragraph_text in section.content.split("\n\n"):
                text = paragraph_text.strip()
                if not text:
                    continue
                p = doc.add_paragraph(text)
                p.style = doc.styles["Normal"]

        # Tables
        for idx, (table_data, caption_text) in enumerate(
            zip(section.tables, section.table_captions or [""] * len(section.tables))
        ):
            if table_data:
                entry = self._captions.next_table(caption_text)
                self._add_table(table_data, entry)

        # Figures
        for idx, (fig_path, caption_text) in enumerate(
            zip(section.figures, section.figure_captions or [""] * len(section.figures))
        ):
            if fig_path and Path(fig_path).exists():
                entry = self._captions.next_figure(caption_text, fig_path)
                self._add_figure(fig_path, entry)

    def _add_figure(self, image_path: str, entry: CaptionEntry) -> None:
        doc = self._doc
        # Center paragraph for image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        try:
            run = p_img.add_run()
            # Fit image to text width (~16cm for A4 with margins)
            run.add_picture(image_path, width=Cm(14))
        except Exception as exc:
            logger.warning(f"Image insert failed for {image_path}: {exc}")
            p_img.add_run(f"[Изображение: {Path(image_path).name}]")

        # Caption paragraph
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        run = p_cap.add_run(entry.caption)
        run.font.name = self._cfg.font_name
        run.font.size = Pt(12)
        run.italic = True

    def _add_table(self, data: list[list[str]], entry: CaptionEntry) -> None:
        doc = self._doc
        if not data:
            return

        # Caption above table (ГОСТ style)
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_cap.paragraph_format.space_after = Pt(3)
        run = p_cap.add_run(entry.caption)
        run.font.name = self._cfg.font_name
        run.font.size = Pt(12)
        run.italic = True

        rows = len(data)
        cols = max(len(row) for row in data) if data else 1

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        for r, row in enumerate(data):
            for c, cell_text in enumerate(row):
                if c < cols:
                    cell = table.cell(r, c)
                    cell.text = str(cell_text)
                    for para in cell.paragraphs:
                        para.style = doc.styles["Normal"]
                        for run in para.runs:
                            run.font.size = Pt(12)
                            run.font.name = self._cfg.font_name
                    # Bold header row
                    if r == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

        doc.add_paragraph()  # Space after table

    def add_paragraph(self, text: str, bold: bool = False, align: str = "left") -> None:
        doc = self._doc
        p = doc.add_paragraph()
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(text)
        run.font.name = self._cfg.font_name
        run.font.size = Pt(self._cfg.font_size_body)
        run.bold = bold

    def add_conclusion(self, text: str) -> None:
        self._doc.add_heading("Вывод", level=1)
        p = self._doc.add_paragraph(text)
        p.style = self._doc.styles["Normal"]

    def add_page_break(self) -> None:
        self._doc.add_page_break()

    def save(self, output_path: Path | str) -> Path:
        if self._doc is None:
            raise RuntimeError("Document not initialized — call new_document() first")
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(output_path))
        logger.info(f"Report saved: {output_path}")
        return output_path

    def get_document(self) -> Any:
        return self._doc


# Module-level singleton
_builder: DocxBuilder | None = None


def get_docx_builder() -> DocxBuilder:
    global _builder
    if _builder is None:
        _builder = DocxBuilder()
    return _builder
