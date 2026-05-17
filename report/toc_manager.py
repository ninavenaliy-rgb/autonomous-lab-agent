"""
Table of Contents manager for DOCX reports.
Builds and updates TOC with correct page references.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any

from loguru import logger

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class TOCEntry:
    level: int      # 1 = chapter, 2 = section, 3 = subsection
    title: str
    page_estimate: int = 0
    anchor: str = ""


class TOCManager:
    """
    Builds and inserts an automatic Word Table of Contents.
    Uses Word's built-in TOC field (TOC \\o "1-3") for auto-update on open.
    """

    def __init__(self) -> None:
        self._entries: list[TOCEntry] = []

    def register_heading(self, title: str, level: int, page: int = 0) -> None:
        self._entries.append(TOCEntry(level=level, title=title, page_estimate=page))

    def insert_toc_field(self, document: Any) -> None:
        """
        Insert a Word TOC field code at the current cursor position.
        Word will populate it when the document is opened and updated.
        """
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available — skipping TOC insert")
            return
        try:
            # Add TOC heading
            toc_heading = document.add_paragraph("Оглавление", style="Heading 1")

            # Add the TOC paragraph with TOC field
            toc_para = document.add_paragraph()
            self._add_toc_field(toc_para)

            # Page break after TOC
            document.add_page_break()
            logger.info("TOC field inserted")
        except Exception as exc:
            logger.error(f"TOC insert failed: {exc}")

    def _add_toc_field(self, paragraph: Any) -> None:
        """Inject the TOC field XML into a paragraph."""
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = r' TOC \o "1-3" \h \z \u '

        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_sep)
        run._r.append(fld_char_end)

    def build_text_toc(self) -> str:
        """Generate a plain-text representation of the TOC (for inspection)."""
        lines = ["ОГЛАВЛЕНИЕ", "=" * 60]
        for entry in self._entries:
            indent = "  " * (entry.level - 1)
            page_str = f"  {entry.page_estimate}" if entry.page_estimate else ""
            lines.append(f"{indent}{entry.title}{page_str}")
        return "\n".join(lines)

    def update_toc_in_document(self, doc_path: str) -> bool:
        """
        Open document in Word and force TOC update via COM.
        Requires Windows + Word installed.
        """
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            # Update all fields (including TOC)
            doc.Fields.Update()
            for toc in doc.TablesOfContents:
                toc.Update()
            doc.Save()
            doc.Close()
            word.Quit()
            logger.info(f"TOC updated in: {doc_path}")
            return True
        except Exception as exc:
            logger.warning(f"TOC COM update failed: {exc}")
            return False

    def inject_update_macro(self, document: Any) -> None:
        """
        Embed a VBA macro in document to auto-update TOC on open.
        Uses word's Document_Open event.
        """
        # This approach embeds a macro via Document.VBA code
        # Only works if macros are enabled; note as best-effort
        macro_code = """
Private Sub Document_Open()
    ActiveDocument.TablesOfContents(1).Update
    ActiveDocument.Fields.Update
End Sub
"""
        try:
            vba_module = document.part.document.element
            # python-docx doesn't expose VBA directly; noted as enhancement
            logger.debug("VBA auto-update macro injection requires COM — skipping")
        except Exception:
            pass
