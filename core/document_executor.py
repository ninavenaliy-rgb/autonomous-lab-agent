"""
Headless Document Executor.
Creates and modifies Word/Excel documents directly via python-docx and openpyxl.
No GUI, no Windows, no Office required. Works on any Linux/Mac/Windows.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

from agents.parser_agent import AtomicStep, ParsedTask
from core.config import get_config


@dataclass
class DocState:
    """Current state of a Word document being edited."""
    path: Path
    doc: Any = None                  # docx.Document
    current_paragraph: Any = None
    cursor_section: int = 0

    def open(self) -> None:
        if self.path.exists():
            self.doc = Document(str(self.path))
        else:
            self.doc = Document()

    def save(self) -> None:
        if self.doc:
            self.doc.save(str(self.path))


@dataclass
class SheetState:
    """Current state of an Excel workbook being edited."""
    path: Path
    wb: Any = None                   # openpyxl.Workbook
    ws: Any = None                   # active worksheet
    current_cell: str = "A1"

    def open(self) -> None:
        if self.path.exists():
            self.wb = openpyxl.load_workbook(str(self.path))
            self.ws = self.wb.active
        else:
            self.wb = openpyxl.Workbook()
            self.ws = self.wb.active

    def save(self) -> None:
        if self.wb:
            self.wb.save(str(self.path))

    def cell(self, address: str) -> Any:
        return self.ws[address]


class HeadlessWordExecutor:
    """
    Executes Word lab tasks via python-docx API.
    Covers all common formatting/structuring operations.
    """

    def __init__(self, output_dir: Path) -> None:
        self._dir = output_dir
        self._dir.mkdir(parents=True, exist_ok=True)
        self._state: DocState | None = None

    def execute(self, action: str, **kwargs) -> bool:
        """Dispatch action. Returns True on success."""
        handler = getattr(self, f"_act_{action}", self._act_unknown)
        try:
            handler(**kwargs)
            return True
        except Exception as exc:
            logger.warning(f"Word action '{action}' failed: {exc}")
            return False

    # ── File operations ───────────────────────────────────────────────────────

    def _act_open_word(self, file_path: str = "", **_) -> None:
        path = Path(file_path) if file_path else self._dir / "document.docx"
        self._state = DocState(path=path)
        self._state.open()
        logger.info(f"Word document: {path}")

    def _act_open_file(self, target: str = "", **_) -> None:
        self._act_open_word(file_path=target)

    def _act_save_file(self, **_) -> None:
        if self._state:
            self._state.save()
            logger.info(f"Saved: {self._state.path}")

    def _act_save_as(self, target: str = "", **_) -> None:
        if self._state and target:
            self._state.path = Path(target)
        self._act_save_file()

    def _act_close_file(self, **_) -> None:
        self._act_save_file()
        self._state = None

    # ── Text input ────────────────────────────────────────────────────────────

    def _act_type(self, value: str = "", target: str = "", **_) -> None:
        doc = self._get_doc()
        # Split by newlines — each becomes a paragraph
        lines = value.split("\n")
        for i, line in enumerate(lines):
            if i == 0 and self._state.current_paragraph is not None:
                self._state.current_paragraph.add_run(line)
            else:
                p = doc.add_paragraph(line)
                self._apply_normal_style(p)
                self._state.current_paragraph = p

    def _act_hotkey(self, value: str = "", **_) -> None:
        # Map common hotkeys to document actions
        hotkey_map = {
            "enter": self._newline,
            "ctrl+enter": self._page_break,
        }
        fn = hotkey_map.get(value.lower())
        if fn:
            fn()

    def _newline(self) -> None:
        doc = self._get_doc()
        p = doc.add_paragraph()
        self._state.current_paragraph = p

    def _page_break(self) -> None:
        doc = self._get_doc()
        doc.add_page_break()
        self._state.current_paragraph = None

    def _act_insert_page_break(self, **_) -> None:
        self._page_break()

    # ── Formatting ────────────────────────────────────────────────────────────

    def _act_apply_style(self, value: str = "", **_) -> None:
        doc = self._get_doc()
        style_map = {
            "заголовок 1": "Heading 1",
            "заголовок 2": "Heading 2",
            "заголовок 3": "Heading 3",
            "обычный": "Normal",
            "heading 1": "Heading 1",
            "heading 2": "Heading 2",
        }
        docx_style = style_map.get(value.lower(), value)
        try:
            p = doc.add_paragraph(style=docx_style)
            self._state.current_paragraph = p
        except Exception:
            # Style not in template — add as heading manually
            level = 1
            m = re.search(r"(\d)", value)
            if m:
                level = int(m.group(1))
            p = doc.add_heading("", level=level)
            self._state.current_paragraph = p

    def _act_format_text(self, value: str = "", **_) -> None:
        p = self._state.current_paragraph if self._state else None
        if not p:
            return
        for run in p.runs:
            if "bold" in value.lower():
                run.bold = True
            if "italic" in value.lower():
                run.italic = True
            if "underline" in value.lower():
                run.underline = True

    def _act_bold(self, **_) -> None:
        if self._state and self._state.current_paragraph:
            for run in self._state.current_paragraph.runs:
                run.bold = not run.bold  # toggle

    def _act_italic(self, **_) -> None:
        if self._state and self._state.current_paragraph:
            for run in self._state.current_paragraph.runs:
                run.italic = not run.italic

    def _act_underline(self, **_) -> None:
        if self._state and self._state.current_paragraph:
            for run in self._state.current_paragraph.runs:
                run.underline = not run.underline

    def _act_set_font(self, value: str = "", **_) -> None:
        if not (self._state and self._state.current_paragraph):
            return
        for run in self._state.current_paragraph.runs:
            run.font.name = value

    def _act_set_font_size(self, value: str = "", **_) -> None:
        if not (self._state and self._state.current_paragraph):
            return
        try:
            size = int(float(value))
            for run in self._state.current_paragraph.runs:
                run.font.size = Pt(size)
        except ValueError:
            pass

    def _act_align(self, value: str = "left", **_) -> None:
        p = self._state.current_paragraph if self._state else None
        if not p:
            return
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "по центру": WD_ALIGN_PARAGRAPH.CENTER,
            "по ширине": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        p.alignment = align_map.get(value.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    def _act_set_line_spacing(self, value: str = "1.5", **_) -> None:
        p = self._state.current_paragraph if self._state else None
        if not p:
            return
        try:
            spacing = float(value)
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = Pt(14 * spacing)
        except ValueError:
            pass

    # ── Insertion ─────────────────────────────────────────────────────────────

    def _act_insert_table(self, rows: int = 3, cols: int = 3,
                           value: str = "", target: str = "", **_) -> None:
        doc = self._get_doc()
        try:
            r = int(rows) if rows else int(value or 3)
            c = int(cols) if cols else int(target or 3)
        except (ValueError, TypeError):
            r, c = 3, 3
        table = doc.add_table(rows=r, cols=c)
        table.style = "Table Grid"
        # Bold header row
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        logger.info(f"Table inserted: {r}×{c}")

    def _act_insert_toc(self, **_) -> None:
        from report.toc_manager import TOCManager
        toc = TOCManager()
        doc = self._get_doc()
        toc.insert_toc_field(doc)

    def _act_insert_image(self, target: str = "", value: str = "", **_) -> None:
        path = Path(target or value)
        if not path.exists():
            logger.warning(f"Image not found: {path}")
            return
        doc = self._get_doc()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(14))

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _get_doc(self):
        if not self._state:
            self._act_open_word()
        if not self._state.doc:
            self._state.open()
        return self._state.doc

    def _apply_normal_style(self, paragraph) -> None:
        try:
            paragraph.style = paragraph._p.getparent().getparent().styles["Normal"]
        except Exception:
            pass

    def _act_wait_ms(self, **_) -> None:
        pass  # No-op in headless mode

    def _act_select_text(self, **_) -> None:
        pass  # No-op — selection is implicit

    def _act_select_all(self, **_) -> None:
        pass

    def _act_focus_window(self, **_) -> None:
        pass

    def _act_unknown(self, **kwargs) -> None:
        pass

    @property
    def document_path(self) -> Path | None:
        return self._state.path if self._state else None


class HeadlessExcelExecutor:
    """
    Executes Excel lab tasks via openpyxl API.
    """

    def __init__(self, output_dir: Path) -> None:
        self._dir = output_dir
        self._dir.mkdir(parents=True, exist_ok=True)
        self._state: SheetState | None = None

    def execute(self, action: str, **kwargs) -> bool:
        handler = getattr(self, f"_act_{action}", self._act_unknown)
        try:
            handler(**kwargs)
            return True
        except Exception as exc:
            logger.warning(f"Excel action '{action}' failed: {exc}")
            return False

    def _act_open_excel(self, file_path: str = "", **_) -> None:
        path = Path(file_path) if file_path else self._dir / "workbook.xlsx"
        self._state = SheetState(path=path)
        self._state.open()
        logger.info(f"Excel workbook: {path}")

    def _act_open_file(self, target: str = "", **_) -> None:
        self._act_open_excel(file_path=target)

    def _act_save_file(self, **_) -> None:
        if self._state:
            self._state.save()

    def _act_click_cell(self, target: str = "", **_) -> None:
        if self._state and target:
            self._state.current_cell = target.upper()

    def _act_type(self, value: str = "", target: str = "", **_) -> None:
        ws = self._get_ws()
        cell_addr = target.upper() if target else self._state.current_cell
        ws[cell_addr] = value
        # Auto-advance to next row
        m = re.match(r"([A-Z]+)(\d+)", cell_addr)
        if m:
            col, row = m.group(1), int(m.group(2))
            self._state.current_cell = f"{col}{row + 1}"

    def _act_enter_formula(self, target: str = "", value: str = "", **_) -> None:
        ws = self._get_ws()
        addr = target.upper() if target else self._state.current_cell
        formula = value if value.startswith("=") else f"={value}"
        ws[addr] = formula
        logger.info(f"Formula: {addr} = {formula}")

    def _act_auto_sum(self, target: str = "", **_) -> None:
        ws = self._get_ws()
        addr = target.upper() if target else self._state.current_cell
        m = re.match(r"([A-Z]+)(\d+)", addr)
        if m:
            col, row = m.group(1), int(m.group(2))
            if row > 1:
                ws[addr] = f"=SUM({col}1:{col}{row - 1})"

    def _act_format_cells(self, value: str = "", **_) -> None:
        ws = self._get_ws()
        # Bold the current row's header
        addr = self._state.current_cell
        ws[addr].font = Font(bold=True)

    def _act_create_chart(
        self,
        chart_type: str = "column",
        data_range: str = "",
        target: str = "",
        **_,
    ) -> None:
        ws = self._get_ws()
        chart_map = {
            "column": BarChart,
            "bar": BarChart,
            "line": LineChart,
            "pie": PieChart,
        }
        ChartClass = chart_map.get(chart_type.lower(), BarChart)
        chart = ChartClass()
        chart.title = "Диаграмма"
        chart.style = 10

        # Auto-detect data range if not specified
        if data_range:
            parts = data_range.split(":")
            if len(parts) == 2:
                try:
                    data = Reference(ws, range_string=f"{ws.title}!{data_range}")
                    chart.add_data(data, titles_from_data=True)
                except Exception:
                    pass

        anchor = target or "E2"
        ws.add_chart(chart, anchor)
        logger.info(f"Chart ({chart_type}) added at {anchor}")

    def _act_excel_create_chart(self, **kwargs) -> None:
        self._act_create_chart(**kwargs)

    def _act_excel_click_cell(self, target: str = "", **_) -> None:
        self._act_click_cell(target=target)

    def _act_excel_enter_formula(self, **kwargs) -> None:
        self._act_enter_formula(**kwargs)

    def _act_excel_auto_sum(self, **kwargs) -> None:
        self._act_auto_sum(**kwargs)

    def _act_excel_format_cells(self, **kwargs) -> None:
        self._act_format_cells(**kwargs)

    def _get_ws(self):
        if not self._state:
            self._act_open_excel()
        if not self._state.ws:
            self._state.open()
        return self._state.ws

    def _act_wait_ms(self, **_) -> None:
        pass

    def _act_focus_window(self, **_) -> None:
        pass

    def _act_unknown(self, **kwargs) -> None:
        pass

    @property
    def document_path(self) -> Path | None:
        return self._state.path if self._state else None


class HeadlessExecutor:
    """
    Unified headless executor. Routes to Word or Excel handler by task.
    Drop-in replacement for GUIAgent when running without Windows/Office.
    """

    def __init__(self) -> None:
        cfg = get_config()
        out = cfg.storage.report_output_dir / "documents"
        out.mkdir(parents=True, exist_ok=True)
        self._word = HeadlessWordExecutor(out)
        self._excel = HeadlessExcelExecutor(out)
        self._current_app = "word"

    def execute(self, action: str, **kwargs) -> Any:
        """Execute action. Returns ActionResult-compatible object."""
        import platform
        app = kwargs.get("application", self._current_app)

        # Route to correct executor
        if action in ("open_excel", "click_cell", "enter_formula",
                       "create_chart", "auto_sum", "format_cells",
                       "excel_click_cell", "excel_enter_formula",
                       "excel_create_chart", "excel_auto_sum"):
            app = "excel"
        if action == "open_word":
            app = "word"

        self._current_app = app
        executor = self._excel if app == "excel" else self._word

        success = executor.execute(action, **kwargs)

        # After save, take a real screenshot from Word/Excel on Mac
        screenshot_path = None
        if success and action in ("save_file", "save_as") and platform.system() == "Darwin":
            doc_path = self.get_document_path(app)
            if doc_path and doc_path.exists():
                screenshot_path = self._capture_mac_screenshot(doc_path, app)

        from agents.gui_agent import ActionResult, ActionStatus
        result = ActionResult(
            status=ActionStatus.SUCCESS if success else ActionStatus.FAILED,
            action=action,
            target=kwargs.get("target", ""),
        )
        # Attach screenshot path as extra attribute if captured
        if screenshot_path:
            object.__setattr__(result, "_screenshot_path", screenshot_path)
        return result

    def _capture_mac_screenshot(self, doc_path: Path, app: str) -> Path | None:
        """Open document in Word/Excel on Mac, take screenshot, close."""
        try:
            from vision.mac_word_capture import take_single_word_screenshot
            cfg = get_config()
            shot_dir = cfg.storage.screenshot_dir
            shot_dir.mkdir(parents=True, exist_ok=True)
            shot_path = shot_dir / f"word_{doc_path.stem}_{int(__import__('time').time())}.png"
            result = take_single_word_screenshot(doc_path, shot_path)
            if result:
                logger.info(f"Mac Word screenshot: {result.name}")
                return result
        except Exception as exc:
            logger.debug(f"Mac screenshot skipped: {exc}")
        return None

    def get_document_path(self, app: str = "word") -> Path | None:
        if app == "excel":
            return self._excel.document_path
        return self._word.document_path
