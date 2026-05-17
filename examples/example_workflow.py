"""
Example workflow: complete a Word formatting lab autonomously.
Demonstrates how to use the agent programmatically.

Run with:
    python examples/example_workflow.py
Or via CLI:
    python app/main.py run examples/sample_lab.docx --dry-run
"""

from __future__ import annotations

import asyncio
import sys
from pathlib import Path

# Ensure project root is importable
sys.path.insert(0, str(Path(__file__).parent.parent))

from loguru import logger
from agents.parser_agent import (
    ParsedTask, AtomicStep, ParsedMethodology, ValidationRule
)
from agents.report_agent import ReportMeta
from core.orchestrator import Orchestrator
from core.config import get_config


# ─────────────────────────────────────────────────────────────────────────────
# Example 1: Run with a real methodology file
# ─────────────────────────────────────────────────────────────────────────────

async def run_with_file(methodology_path: str, output_path: str | None = None):
    """
    Run the full autonomous pipeline on a .docx or .pdf methodology.
    """
    cfg = get_config()
    cfg.ensure_dirs()

    path = Path(methodology_path)
    if not path.exists():
        logger.error(f"File not found: {path}")
        return

    meta = ReportMeta(
        title=path.stem,
        student="Иванов И.И.",
        group="ИС-42",
        teacher="Петрова А.В.",
        university="Технический Университет",
        department="Кафедра Информатики",
        lab_number="1",
    )

    orchestrator = Orchestrator()
    logger.info(f"Starting autonomous execution: {path.name}")
    report = await orchestrator.run(
        methodology_path=path,
        output_path=Path(output_path) if output_path else None,
        resume=True,
        report_meta=meta,
    )
    logger.info(f"Report: {report}")
    return report


# ─────────────────────────────────────────────────────────────────────────────
# Example 2: Inject a manually-defined methodology (no file needed)
# ─────────────────────────────────────────────────────────────────────────────

EXAMPLE_WORD_METHODOLOGY = ParsedMethodology(
    title="Лабораторная работа №1: Основы работы в Microsoft Word",
    document_path="manual",
    language="ru",
    global_context=(
        "В данной лабораторной работе изучаются основы форматирования "
        "текстовых документов в Microsoft Word 365."
    ),
    tasks=[
        ParsedTask(
            task_id="word_1_1",
            application="word",
            title="Создание и форматирование заголовка",
            description="Создайте заголовок документа и примените стиль Заголовок 1",
            section="1. Создание документа",
            order_index=0,
            requires_screenshot=True,
            expected_result="Заголовок отформатирован стилем Заголовок 1",
            steps=[
                AtomicStep(
                    step_id="s1",
                    action="open_word",
                    target="",
                    description="Запустить Microsoft Word",
                    expected_outcome="Word открыт",
                ),
                AtomicStep(
                    step_id="s2",
                    action="type",
                    value="Лабораторная работа №1",
                    description="Ввести заголовок",
                    expected_outcome="Текст введён",
                ),
                AtomicStep(
                    step_id="s3",
                    action="select_text",
                    target="",
                    value="line",
                    description="Выделить строку",
                    expected_outcome="Строка выделена",
                ),
                AtomicStep(
                    step_id="s4",
                    action="apply_style",
                    value="Заголовок 1",
                    description="Применить стиль Заголовок 1",
                    expected_outcome="Стиль применён",
                ),
                AtomicStep(
                    step_id="s5",
                    action="align",
                    value="center",
                    description="Выровнять по центру",
                    expected_outcome="Текст по центру",
                ),
            ],
            validation_rules=[
                ValidationRule(
                    rule_type="text_contains",
                    target="Лабораторная работа",
                    expected_value="Лабораторная работа №1",
                    description="Заголовок присутствует в документе",
                ),
            ],
        ),
        ParsedTask(
            task_id="word_1_2",
            application="word",
            title="Вставка таблицы",
            description="Вставьте таблицу 3×3 с данными",
            section="2. Работа с таблицами",
            order_index=1,
            requires_screenshot=True,
            expected_result="Таблица 3×3 вставлена и заполнена",
            steps=[
                AtomicStep(
                    step_id="s1",
                    action="press_key",
                    value="ctrl+end",
                    description="Перейти в конец документа",
                    expected_outcome="Курсор в конце",
                ),
                AtomicStep(
                    step_id="s2",
                    action="hotkey",
                    value="enter",
                    description="Новый абзац",
                    expected_outcome="",
                ),
                AtomicStep(
                    step_id="s3",
                    action="insert_table",
                    value="3",
                    target="3",
                    description="Вставить таблицу 3×3",
                    expected_outcome="Таблица появилась",
                ),
                AtomicStep(
                    step_id="s4",
                    action="type",
                    value="Наименование\tЗначение\tПримечание",
                    description="Заполнить заголовок таблицы",
                    expected_outcome="Заголовок введён",
                ),
                AtomicStep(
                    step_id="s5",
                    action="save_file",
                    description="Сохранить документ",
                    expected_outcome="Документ сохранён",
                ),
            ],
        ),
        ParsedTask(
            task_id="word_1_3",
            application="word",
            title="Форматирование текста",
            description="Введите абзац текста с форматированием",
            section="3. Форматирование",
            order_index=2,
            requires_screenshot=True,
            expected_result="Абзац с полужирным и курсивным текстом",
            steps=[
                AtomicStep(
                    step_id="s1",
                    action="hotkey",
                    value="ctrl+end",
                    description="Конец документа",
                ),
                AtomicStep(
                    step_id="s2",
                    action="hotkey",
                    value="enter",
                    description="Новый абзац",
                ),
                AtomicStep(
                    step_id="s3",
                    action="type",
                    value="Это обычный текст. ",
                    description="Ввод обычного текста",
                ),
                AtomicStep(
                    step_id="s4",
                    action="bold",
                    description="Включить полужирный",
                    expected_outcome="Полужирный включён",
                ),
                AtomicStep(
                    step_id="s5",
                    action="type",
                    value="Это полужирный текст. ",
                    description="Ввод полужирного",
                ),
                AtomicStep(
                    step_id="s6",
                    action="bold",
                    description="Выключить полужирный",
                ),
                AtomicStep(
                    step_id="s7",
                    action="italic",
                    description="Включить курсив",
                ),
                AtomicStep(
                    step_id="s8",
                    action="type",
                    value="Это курсив.",
                    description="Ввод курсива",
                ),
                AtomicStep(
                    step_id="s9",
                    action="italic",
                    description="Выключить курсив",
                ),
                AtomicStep(
                    step_id="s10",
                    action="save_file",
                    description="Сохранить",
                    expected_outcome="Сохранено",
                ),
            ],
        ),
    ],
)


EXAMPLE_EXCEL_METHODOLOGY = ParsedMethodology(
    title="Лабораторная работа №2: Работа с Microsoft Excel",
    document_path="manual",
    language="ru",
    global_context="Изучение формул и диаграмм в Microsoft Excel 365.",
    tasks=[
        ParsedTask(
            task_id="excel_1_1",
            application="excel",
            title="Создание таблицы с формулами",
            description="Создайте таблицу успеваемости с формулами SUM и AVERAGE",
            order_index=0,
            requires_screenshot=True,
            expected_result="Таблица с формулами заполнена",
            steps=[
                AtomicStep(step_id="s1", action="open_excel", description="Открыть Excel"),
                AtomicStep(
                    step_id="s2", action="click_cell", target="A1",
                    description="Выбрать ячейку A1"
                ),
                AtomicStep(
                    step_id="s3", action="type", value="Студент\tОценка 1\tОценка 2\tСредняя",
                    description="Заголовки"
                ),
                AtomicStep(
                    step_id="s4", action="click_cell", target="A2",
                    description="Следующая строка"
                ),
                AtomicStep(
                    step_id="s5", action="type", value="Иванов\t4\t5",
                    description="Данные строки 1"
                ),
                AtomicStep(
                    step_id="s6", action="click_cell", target="D2",
                    description="Ячейка для формулы"
                ),
                AtomicStep(
                    step_id="s7", action="enter_formula",
                    target="D2", value="=AVERAGE(B2:C2)",
                    description="Формула среднего"
                ),
                AtomicStep(
                    step_id="s8", action="save_file",
                    description="Сохранить"
                ),
            ],
        ),
    ],
)


# ─────────────────────────────────────────────────────────────────────────────
# Example 3: Generate a sample methodology DOCX for testing
# ─────────────────────────────────────────────────────────────────────────────

def create_sample_methodology(output_path: Path | None = None) -> Path:
    """Create a sample .docx methodology file for testing the parser."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        logger.error("python-docx not installed")
        return Path("sample_lab.docx")

    doc = Document()
    doc.add_heading("Лабораторная работа №1", level=0)
    doc.add_heading("Работа с Microsoft Word", level=1)

    doc.add_paragraph(
        "Цель работы: освоить основные приёмы работы с текстовым редактором "
        "Microsoft Word 365."
    )

    doc.add_heading("Задание 1. Создание и форматирование документа", level=2)
    p = doc.add_paragraph(
        "Создайте новый документ Microsoft Word и выполните следующие действия:"
    )
    steps_text = [
        "1. Откройте Microsoft Word и создайте новый пустой документ.",
        "2. Введите заголовок «Лабораторная работа №1» и примените стиль «Заголовок 1».",
        "3. Установите шрифт Times New Roman, размер 14 пт.",
        "4. Выровняйте заголовок по центру страницы.",
        "5. Сохраните документ с именем Lab1.docx.",
    ]
    for step in steps_text:
        doc.add_paragraph(step)

    doc.add_heading("Задание 2. Работа с таблицами", level=2)
    doc.add_paragraph(
        "Создайте таблицу для хранения данных о студентах:"
    )
    steps_table = [
        "1. Вставьте таблицу размером 3×4.",
        "2. Заполните заголовки: Фамилия, Имя, Группа, Оценка.",
        "3. Введите данные трёх студентов.",
        "4. Отформатируйте заголовочную строку жирным шрифтом.",
        "5. Сделайте снимок экрана (скриншот) готовой таблицы.",
    ]
    for step in steps_table:
        doc.add_paragraph(step)

    doc.add_heading("Задание 3. Вставка рисунка", level=2)
    doc.add_paragraph(
        "1. Вставьте любое изображение из файла.\n"
        "2. Добавьте подпись к рисунку.\n"
        "3. Обновите оглавление документа."
    )

    output = output_path or (Path(__file__).parent / "sample_lab.docx")
    doc.save(str(output))
    logger.info(f"Sample methodology created: {output}")
    return output


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Example workflow runner")
    parser.add_argument("--file", help="Methodology file path")
    parser.add_argument("--create-sample", action="store_true",
                        help="Create sample_lab.docx for testing")
    parser.add_argument("--dry-run", action="store_true",
                        help="Parse only, don't execute GUI")
    parser.add_argument("--output", help="Output report path")
    args = parser.parse_args()

    # Setup
    cfg = get_config()
    cfg.ensure_dirs()
    logger.remove()
    logger.add(sys.stderr, level="INFO", colorize=True)

    if args.create_sample:
        sample = create_sample_methodology()
        print(f"Created: {sample}")
        sys.exit(0)

    if args.file:
        asyncio.run(run_with_file(args.file, args.output))
    else:
        # Demo: show parsed structure of the built-in example
        print("=== Example Word Methodology ===")
        for task in EXAMPLE_WORD_METHODOLOGY.tasks:
            print(f"\nTask: {task.task_id} — {task.title}")
            for step in task.steps:
                print(f"  [{step.step_id}] {step.action}: {step.description}")
        print("\nRun with: python examples/example_workflow.py --create-sample")
        print("Then:     python app/main.py run examples/sample_lab.docx")
