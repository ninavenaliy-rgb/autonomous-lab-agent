"""
Instruction Parser Agent.
Reads DOCX/PDF methodology files and extracts an executable task graph via LLM.
Supports Russian-language documents and educational lab instruction formats.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from loguru import logger
from pydantic import BaseModel, Field
from tenacity import retry, stop_after_attempt, wait_exponential

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

from core.config import get_config, LLMProvider


# ─────────────────────────────────────────────────────────────────────────────
# Data Models
# ─────────────────────────────────────────────────────────────────────────────

class ValidationRule(BaseModel):
    rule_type: str  # "element_exists" | "text_contains" | "style_applied" | "screenshot"
    target: str = ""
    expected_value: str = ""
    description: str = ""


class AtomicStep(BaseModel):
    step_id: str
    action: str           # "click" | "type" | "hotkey" | "menu" | "insert_image" | "format" | ...
    target: str = ""      # element name, menu path, text, etc.
    value: str = ""       # text to type, style name, etc.
    description: str = ""
    expected_outcome: str = ""
    optional: bool = False


class ParsedTask(BaseModel):
    task_id: str
    application: str      # "word" | "excel" | "any"
    title: str
    description: str
    section: str = ""     # original document section
    steps: list[AtomicStep] = Field(default_factory=list)
    expected_result: str = ""
    validation_rules: list[ValidationRule] = Field(default_factory=list)
    requires_screenshot: bool = True
    order_index: int = 0


class ParsedMethodology(BaseModel):
    title: str
    document_path: str
    tasks: list[ParsedTask] = Field(default_factory=list)
    global_context: str = ""
    language: str = "ru"
    total_sections: int = 0

    def get_word_tasks(self) -> list[ParsedTask]:
        return [t for t in self.tasks if t.application == "word"]

    def get_excel_tasks(self) -> list[ParsedTask]:
        return [t for t in self.tasks if t.application == "excel"]


# ─────────────────────────────────────────────────────────────────────────────
# Document Readers
# ─────────────────────────────────────────────────────────────────────────────

class DocumentReader:
    """Extracts raw text from DOCX and PDF files, preserving section structure."""

    def read(self, path: Path) -> tuple[str, dict[str, str]]:
        """
        Returns (full_text, sections_dict).
        sections_dict: {heading_text: section_content}
        """
        suffix = path.suffix.lower()
        if suffix == ".docx":
            return self._read_docx(path)
        elif suffix == ".pdf":
            return self._read_pdf(path)
        else:
            raise ValueError(f"Unsupported format: {suffix}")

    def _read_docx(self, path: Path) -> tuple[str, dict[str, str]]:
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        doc = DocxDocument(str(path))
        sections: dict[str, str] = {}
        current_heading = "Introduction"
        current_paragraphs: list[str] = []
        full_lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            full_lines.append(text)
            style_name = para.style.name.lower() if para.style else ""
            if "heading" in style_name or "заголовок" in style_name or self._is_heading(text):
                if current_paragraphs:
                    sections[current_heading] = "\n".join(current_paragraphs)
                current_heading = text
                current_paragraphs = []
            else:
                current_paragraphs.append(text)

        if current_paragraphs:
            sections[current_heading] = "\n".join(current_paragraphs)

        # Also extract tables as text
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            full_lines.append("\n".join(rows))

        return "\n".join(full_lines), sections

    def _read_pdf(self, path: Path) -> tuple[str, dict[str, str]]:
        if not PDF_AVAILABLE:
            raise RuntimeError("pdfplumber not installed")
        lines: list[str] = []
        sections: dict[str, str] = {}
        current_heading = "Introduction"
        current_lines: list[str] = []

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    lines.append(line)
                    if self._is_heading(line):
                        if current_lines:
                            sections[current_heading] = "\n".join(current_lines)
                        current_heading = line
                        current_lines = []
                    else:
                        current_lines.append(line)
                # Tables
                for table in page.extract_tables() or []:
                    for row in table:
                        if row:
                            lines.append(" | ".join(str(c or "") for c in row))

        if current_lines:
            sections[current_heading] = "\n".join(current_lines)

        return "\n".join(lines), sections

    def _is_heading(self, text: str) -> bool:
        """Heuristic: line looks like a section heading."""
        if len(text) > 120:
            return False
        patterns = [
            r"^\d+[\.\)]\s+\S",           # "1. Title" or "1) Title"
            r"^\d+\.\d+\s+\S",            # "1.2 Title"
            r"^[А-ЯA-Z][А-ЯA-Z\s]+$",    # All-caps Russian/English
            r"^(Задание|Задача|Упражнение|Практика|Лаборатор|Task|Exercise)\s",
        ]
        for pat in patterns:
            if re.match(pat, text, re.IGNORECASE):
                return True
        return False


# ─────────────────────────────────────────────────────────────────────────────
# LLM Client
# ─────────────────────────────────────────────────────────────────────────────

class LLMClient:
    """Unified LLM client for Anthropic Claude and OpenAI."""

    def __init__(self) -> None:
        cfg = get_config()
        self._provider = cfg.llm.provider
        self._cfg = cfg.llm
        self._client: Any = None
        self._initialize()

    def _initialize(self) -> None:
        if self._provider == LLMProvider.ANTHROPIC and ANTHROPIC_AVAILABLE:
            self._client = anthropic.Anthropic(api_key=self._cfg.anthropic_api_key)
        elif self._provider == LLMProvider.OPENAI and OPENAI_AVAILABLE:
            self._client = openai.OpenAI(api_key=self._cfg.openai_api_key)
        else:
            logger.warning("No LLM client available — parser will use heuristics only")

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=30),
        reraise=True,
    )
    def complete(self, system: str, user: str) -> str:
        if self._client is None:
            raise RuntimeError("LLM client not initialized")

        if self._provider == LLMProvider.ANTHROPIC:
            response = self._client.messages.create(
                model=self._cfg.claude_model,
                max_tokens=self._cfg.max_tokens,
                temperature=self._cfg.temperature,
                system=system,
                messages=[{"role": "user", "content": user}],
            )
            return response.content[0].text

        elif self._provider == LLMProvider.OPENAI:
            response = self._client.chat.completions.create(
                model=self._cfg.openai_model,
                max_tokens=self._cfg.max_tokens,
                temperature=self._cfg.temperature,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
            )
            return response.choices[0].message.content or ""

        return ""


# ─────────────────────────────────────────────────────────────────────────────
# Parser Agent
# ─────────────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert educational lab assistant analyzer.
Your task is to parse laboratory methodology documents (in Russian or English)
and extract a structured list of executable tasks for a desktop automation agent.

The agent can control Microsoft Word and Excel on Windows 11.
Each task must be broken into atomic steps with clear actions.

IMPORTANT RULES:
1. Detect whether each task uses Word or Excel (look for document/table context)
2. Steps must be atomic: one action per step
3. Include validation rules to verify each step succeeded
4. Preserve section numbering from the original document
5. Use action types: click, type, hotkey, menu_navigate, insert_image,
   apply_style, insert_table, format_text, save_file, open_file,
   select_text, scroll, create_chart, enter_formula
6. For Russian UI: use Russian button/menu names in target fields
7. Mark steps that require a screenshot: requires_screenshot=true

Return ONLY valid JSON matching this schema — no prose, no code fences:
{
  "title": "Lab title",
  "language": "ru",
  "global_context": "Brief lab description",
  "tasks": [
    {
      "task_id": "word_1_1",
      "application": "word",
      "title": "Task title",
      "description": "What to do",
      "section": "Original section heading",
      "order_index": 0,
      "requires_screenshot": true,
      "expected_result": "What the document should look like",
      "steps": [
        {
          "step_id": "s1",
          "action": "open_file",
          "target": "path or dialog",
          "value": "",
          "description": "Open Word document",
          "expected_outcome": "Word is open with document",
          "optional": false
        }
      ],
      "validation_rules": [
        {
          "rule_type": "element_exists",
          "target": "document title",
          "expected_value": "",
          "description": "Document title appears"
        }
      ]
    }
  ]
}
"""


class ParserAgent:
    """
    Reads methodology files and produces a ParsedMethodology via LLM analysis.
    Falls back to heuristic parsing if LLM is unavailable.
    """

    def __init__(self) -> None:
        self._reader = DocumentReader()
        self._llm = LLMClient()
        self._chunk_size = 6000  # chars per LLM request to avoid context overflow

    def parse(self, path: Path) -> ParsedMethodology:
        """Parse a methodology file and return structured task graph."""
        if not path.exists():
            raise FileNotFoundError(f"Methodology file not found: {path}")
        logger.info(f"Parsing methodology: {path.name}")

        full_text, sections = self._reader.read(path)
        logger.info(f"Extracted {len(full_text)} chars, {len(sections)} sections")

        if self._llm._client is not None:
            return self._parse_with_llm(path, full_text, sections)
        else:
            return self._parse_heuristic(path, full_text, sections)

    def _parse_with_llm(
        self, path: Path, full_text: str, sections: dict[str, str]
    ) -> ParsedMethodology:
        """Use LLM to extract structured tasks."""
        # Split into chunks to handle long documents
        chunks = self._split_text(full_text, self._chunk_size)
        all_tasks: list[ParsedTask] = []
        global_context = ""
        doc_title = path.stem

        for i, chunk in enumerate(chunks):
            logger.debug(f"LLM parsing chunk {i+1}/{len(chunks)}")
            prompt = f"""Parse this laboratory methodology document chunk and extract all tasks.
Document: {path.name}
Chunk {i+1} of {len(chunks)}:

{chunk}

Extract ALL laboratory tasks/exercises mentioned. For each task, identify:
- The application (word/excel)
- Step-by-step instructions
- Expected results
- Required validations
"""
            try:
                response = self._llm.complete(SYSTEM_PROMPT, prompt)
                parsed = self._extract_json(response)
                if parsed:
                    if i == 0:
                        doc_title = parsed.get("title", path.stem)
                        global_context = parsed.get("global_context", "")
                    raw_tasks = parsed.get("tasks", [])
                    offset = len(all_tasks)
                    for j, rt in enumerate(raw_tasks):
                        rt["order_index"] = offset + j
                        try:
                            task = ParsedTask(**rt)
                            all_tasks.append(task)
                        except Exception as exc:
                            logger.warning(f"Failed to parse task {j}: {exc}")
            except Exception as exc:
                logger.error(f"LLM parsing chunk {i} failed: {exc}")

        if not all_tasks:
            logger.warning("LLM returned no tasks — falling back to heuristics")
            return self._parse_heuristic(path, full_text, sections)

        return ParsedMethodology(
            title=doc_title,
            document_path=str(path),
            tasks=all_tasks,
            global_context=global_context,
            total_sections=len(sections),
        )

    def _parse_heuristic(
        self, path: Path, full_text: str, sections: dict[str, str]
    ) -> ParsedMethodology:
        """Rule-based fallback parser for when LLM is unavailable."""
        logger.info("Using heuristic parser")
        tasks: list[ParsedTask] = []
        order = 0

        for heading, content in sections.items():
            app = self._detect_application(heading + " " + content)
            task_id = f"{app}_{order + 1}"
            steps = self._extract_steps_heuristic(content)
            tasks.append(ParsedTask(
                task_id=task_id,
                application=app,
                title=heading,
                description=content[:500],
                section=heading,
                steps=steps,
                expected_result="",
                requires_screenshot=True,
                order_index=order,
            ))
            order += 1

        return ParsedMethodology(
            title=path.stem,
            document_path=str(path),
            tasks=tasks,
            total_sections=len(sections),
        )

    def _extract_steps_heuristic(self, text: str) -> list[AtomicStep]:
        """Extract numbered steps from text using regex."""
        steps = []
        pattern = re.compile(
            r"^\s*(\d+[\.\)])\s+(.+)$", re.MULTILINE
        )
        for i, match in enumerate(pattern.finditer(text)):
            step_text = match.group(2).strip()
            action = self._infer_action(step_text)
            steps.append(AtomicStep(
                step_id=f"s{i+1}",
                action=action,
                target="",
                value=step_text[:200],
                description=step_text,
                expected_outcome="",
            ))
        return steps

    def _detect_application(self, text: str) -> str:
        """Detect whether content describes Word or Excel tasks."""
        text_lower = text.lower()
        excel_keywords = [
            "excel", "таблиц", "формул", "ячейк", "лист", "диаграмм",
            "spreadsheet", "worksheet", ".xlsx", "сумм", "среднее",
        ]
        word_keywords = [
            "word", "документ", "абзац", "шрифт", "заголовок", "оглавлени",
            "paragraph", "font", ".docx", "форматирован", "стил",
        ]
        excel_score = sum(1 for k in excel_keywords if k in text_lower)
        word_score = sum(1 for k in word_keywords if k in text_lower)
        if excel_score > word_score:
            return "excel"
        if word_score > 0:
            return "word"
        return "word"  # default

    def _infer_action(self, text: str) -> str:
        """Infer action type from instruction text."""
        text_lower = text.lower()
        if any(k in text_lower for k in ["открыт", "open", "запуст"]):
            return "open_file"
        if any(k in text_lower for k in ["введ", "напишит", "введите", "type", "вписат"]):
            return "type"
        if any(k in text_lower for k in ["нажм", "press", "кнопк", "click"]):
            return "click"
        if any(k in text_lower for k in ["выдел", "select", "отмет"]):
            return "select_text"
        if any(k in text_lower for k in ["сохран", "save"]):
            return "save_file"
        if any(k in text_lower for k in ["вставьт", "insert", "добавьт"]):
            return "insert_image"
        if any(k in text_lower for k in ["формат", "format", "шрифт", "font"]):
            return "format_text"
        if any(k in text_lower for k in ["стил", "style"]):
            return "apply_style"
        if any(k in text_lower for k in ["таблиц", "table"]):
            return "insert_table"
        if any(k in text_lower for k in ["диаграмм", "chart", "график"]):
            return "create_chart"
        if any(k in text_lower for k in ["формул", "=", "function"]):
            return "enter_formula"
        return "type"

    def _split_text(self, text: str, chunk_size: int) -> list[str]:
        """Split text into chunks at paragraph boundaries."""
        if len(text) <= chunk_size:
            return [text]
        chunks = []
        paragraphs = text.split("\n")
        current = []
        current_len = 0
        for para in paragraphs:
            if current_len + len(para) > chunk_size and current:
                chunks.append("\n".join(current))
                current = []
                current_len = 0
            current.append(para)
            current_len += len(para)
        if current:
            chunks.append("\n".join(current))
        return chunks

    def _extract_json(self, text: str) -> dict | None:
        """Extract JSON from LLM response, stripping any surrounding prose."""
        # Remove code fences
        text = re.sub(r"```(?:json)?\s*", "", text)
        text = re.sub(r"```\s*$", "", text)
        # Find JSON object
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if not match:
            logger.warning("No JSON found in LLM response")
            return None
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError as exc:
            # Try to fix common issues
            fixed = match.group(0)
            fixed = re.sub(r",\s*}", "}", fixed)
            fixed = re.sub(r",\s*]", "]", fixed)
            try:
                return json.loads(fixed)
            except Exception:
                logger.error(f"JSON parse failed: {exc}")
                return None


# Module-level singleton
_parser: ParserAgent | None = None


def get_parser_agent() -> ParserAgent:
    global _parser
    if _parser is None:
        _parser = ParserAgent()
    return _parser
