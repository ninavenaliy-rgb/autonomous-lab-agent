"""
Document Cloner.
Takes a completed reference .docx and returns an exact copy
with the student name / group substituted throughout.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from loguru import logger

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def detect_student_info(doc) -> dict[str, str]:
    """Scan title page paragraphs for original student name and group."""
    info: dict[str, str] = {"name": "", "group": ""}
    for para in list(doc.paragraphs)[:40]:
        text = para.text.strip()
        if not text:
            continue
        # "Выполнил: Иванова Юлия" / "Студент Иванова Ю.И."
        m = re.search(
            r"(?:Выполнил[аи]?|Студент[ка]?|Автор)[:\s]+([А-ЯЁа-яёA-Za-z][^\n\r,;]{2,40})",
            text,
        )
        if m and not info["name"]:
            info["name"] = m.group(1).strip()
        # "Группа: ИВТ-21" / "Гр. ЭВМ-301"
        m = re.search(
            r"(?:Группа|Гр\.?)[:\s]+([А-ЯЁа-яёA-Z0-9][А-ЯЁа-яёA-Z0-9\-]{1,15})",
            text,
        )
        if m and not info["group"]:
            info["group"] = m.group(1).strip()
    return info


def _replace_in_para(para, old: str, new: str) -> int:
    """Replace old→new inside a paragraph preserving run formatting."""
    count = 0
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    return count


def clone_document(
    reference_path: Path,
    output_path: Path,
    new_name: str = "",
    new_group: str = "",
) -> Path:
    """
    Copy reference_path to output_path and substitute student info.
    Returns output_path.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(reference_path), str(output_path))

    if not new_name and not new_group:
        logger.info("Clone: no substitutions requested, returning raw copy")
        return output_path

    doc = Document(str(output_path))
    original = detect_student_info(doc)
    logger.info(f"Clone: detected original info: {original}")

    # Build substitution pairs (old → new)
    subs: list[tuple[str, str]] = []
    if new_name and original.get("name"):
        subs.append((original["name"], new_name))
    if new_group and original.get("group"):
        subs.append((original["group"], new_group))

    if not subs:
        logger.warning("Clone: no patterns found to substitute — returning raw copy")
        return output_path

    total = 0

    # Body paragraphs
    for para in doc.paragraphs:
        for old, new in subs:
            total += _replace_in_para(para, old, new)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in subs:
                        total += _replace_in_para(para, old, new)

    # Headers and footers
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for para in hdr.paragraphs:
                for old, new in subs:
                    total += _replace_in_para(para, old, new)

    doc.save(str(output_path))
    logger.info(f"Clone saved: {output_path.name} ({total} substitutions)")
    return output_path
